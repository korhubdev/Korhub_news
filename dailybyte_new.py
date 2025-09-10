import streamlit as st
import requests
from bs4 import BeautifulSoup
from datetime import datetime
import time
import os
from docx import Document
from io import BytesIO

# ✅ Streamlit UI 제목
st.markdown("<h3 style='text-align: center; color: black;'>📝 Dailybyte Challenge 기사 수정 및 저장</h3>", unsafe_allow_html=True)

# ✅ 날짜 선택 (오늘 기본값)
selected_date = st.date_input("🗓️ 날짜 선택", datetime.today())

# ✅ 날짜 변수 세팅
today_date = selected_date.strftime("%Y년 %m월 %d일")
file_date = selected_date.strftime('%Y%m%d')
news_date = selected_date.strftime('%y%m%d')
audio_date = selected_date.strftime('%Y-%m-%d')

# ✅ 세션 상태 초기화
if "edit_mode" not in st.session_state:
    st.session_state.edit_mode = False
if "articles" not in st.session_state:
    st.session_state.articles = []

# ✅ 기사 가져오기 (선택 없이 전부 수정 대상으로 설정)
if st.button("🔄 기사 가져오기"):
    with st.spinner("🕒 기사를 불러오는 중..."):
        time.sleep(3)
        url = f"https://www.mydailybyte.com/post/realtimenews{news_date}l"
        headers = {
            "Cache-Control": "no-cache",
            "Pragma": "no-cache",
            "User-Agent": "Mozilla/5.0"
        }
        response = requests.get(url, headers=headers)

        if response.status_code != 200:
            st.warning("⚠️ 해당 날짜의 기사를 불러올 수 없습니다.")
        else:
            soup = BeautifulSoup(response.text, "html.parser")
            articles = soup.find_all('h2')
            extracted_articles = []

            for idx, article in enumerate(articles, start=1):
                title = article.text.strip()
                content_tag = article.find_next_sibling('p')
                if content_tag and content_tag.text.strip():
                    content = content_tag.text.strip()
                    extracted_articles.append({"index": idx, "title": title, "content": content})

            if not extracted_articles:
                st.warning("⚠️ 해당 날짜의 기사가 아직 업로드되지 않았습니다.")
            else:
                st.session_state.articles = extracted_articles
                st.session_state.edit_mode = True  # 👉 바로 수정모드 진입
                st.rerun()

# ✅ 수정 화면
if st.session_state.edit_mode:
    st.write("✏️ **전체 기사 수정:**")
    edited_articles = []

    for article in st.session_state.articles:
        with st.expander(f"📌 {article['title']}"):
            title = st.text_input("제목", article["title"], key=f"title_{article['index']}", disabled=True)
            content = st.text_area("내용 수정", article["content"], key=f"content_{article['index']}")
            edited_articles.append({"title": title, "content": content})

    # ✅ 대본 저장 버튼 (루프 바깥에 있어야 함)
    if st.button("📜 대본 다운로드", key="download_button"):
        doc = Document()
        for article in edited_articles:
            doc.add_paragraph(f"{today_date} 한눈에 보는 오늘의 뉴스입니다.\n")
            doc.add_paragraph(article["content"])
            doc.add_paragraph("지금까지 코럽 뉴스캐스터 (                 ) 였습니다.\n")
            doc.add_paragraph("-" * 50)

        # Word 파일을 메모리에 저장 (로컬에 저장 X)
        output_script = BytesIO()
        doc.save(output_script)
        output_script.seek(0)

        st.success("✅ 대본 Word 파일이 준비되었습니다! 다운로드 버튼을 눌러주세요.")
        st.download_button(
            label="📅 대본 다운로드",
            data=output_script,
            file_name=f"Dailybyte_Script_{file_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download"
        )

